import argparse
import os
import fitz
from google.cloud import vision
from PIL import Image
from io import BytesIO
from docx import Document
from openai import OpenAI

openai_client = OpenAI()

def get_images_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    images = []
    for page_number in range(len(doc)):
        print(f"📸 Converting to image: Page {page_number + 1}/{len(doc)}...")
        page = doc[page_number]
        pix = page.get_pixmap(dpi=300)
        image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(image)
    return images

def image_to_text(image):
    client = vision.ImageAnnotatorClient()

    buffered = BytesIO()
    image.save(buffered, format="PNG")
    image_content = buffered.getvalue()

    image = vision.Image(content=image_content)

    image_context = vision.ImageContext(language_hints=["de"])

    response = client.text_detection(image=image, image_context=image_context)
    if response.error.message:
        print(f"Google Vision API error: {response.error.message}")
    
    return response.text_annotations[0].description if response.text_annotations else ""

def translate_text_llm(text, prompt):
    try:
        completion = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system", 
                    "content": prompt
                },
                {
                    "role": "user",
                    "content": text
                }
            ]
        )
        return completion.choices[0].message.content
    except Exception as e:
        print(f"⚠️ OpenAI API error: {e}")
        return "########################\n" \
                + "OpenAI encountered an error here, check error log and ensure account has funding\n" \
                + "########################\n" + text

def save_to_word(translations, output_path):
    doc = Document()
    for page_number, translation in enumerate(translations, start=1):
        doc.add_heading(f"Page {page_number}", level=3)
        doc.add_paragraph(translation)
    doc.save(output_path)

def main(pdf_path, output_docx):
    # Load GPT system prompt
    script_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(script_dir, "prompt.txt")
    prompt = ""
    with open(file_path, "r", encoding="utf-8") as file:
        prompt = file.read()

    # Get all images from the PDF
    images = get_images_from_pdf(pdf_path)

    # Extract all text from images
    extracted_text_pages = []
    for i in range(len(images)):
        print(f"👁️ Analyzing image and extracting text: Page {i + 1}/{len(images)}...")
        extracted_text_pages.append(image_to_text(images[i]))
    
    translations = []
    for page_number in range(len(extracted_text_pages)):
        print(f"💭 Translating page {page_number + 1}/{len(extracted_text_pages)}...")
        page_translation = ""
        translated_text = translate_text_llm(extracted_text_pages[page_number], prompt)
        if translated_text != "":
            page_translation += translated_text
        translations.append(page_translation)
    save_to_word(translations, output_docx)
    print(f"✅ Success! Translated document saved to {output_docx}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Translate a PDF and save the result in Word format.")
    parser.add_argument("input_pdf", type=str, help="Path to the input PDF file")
    args = parser.parse_args()

    input_pdf = args.input_pdf
    output_dir = "output"
    input_filename = os.path.basename(input_pdf)
    input_name_without_ext = os.path.splitext(input_filename)[0]
    output_word = os.path.join(output_dir, f"{input_name_without_ext}_translated.docx")

    os.makedirs(output_dir, exist_ok=True)

    main(input_pdf, output_word)
